from pypdf import PdfReader
from docx import Document
import os

def extract_title_from_pdf(pdf_path,fallback=None):
    """
    Extracts the title of a PDF file.

    This function attempts the following in order:
    1. Reads the title from PDF metadata.
    2. Uses the first line of text from the first page if metadata is missing.
    3. Uses a fallback value if provided.
    4. Falls back to the filename as the title.

    Parameters:
        pdf_path (str): Path to the PDF file.
        fallback (str, optional): Fallback title if no metadata or content is found.

    Returns:
        str: The extracted title.
    """
    reader = PdfReader(pdf_path)
    
    # Try to get the title from metadata
    if reader.metadata and reader.metadata.title:
        return reader.metadata.title.strip()
    
    # If no metadata title, infer from the first page
    if len(reader.pages) > 0:
        first_page_text = reader.pages[0].extract_text().strip()
        if first_page_text:
            return first_page_text.split("\n")[0]  # Assume first line is the title
    if fallback:
        return fallback
    return os.path.basename(pdf_path)

def extract_title_from_docx(docx_path,fallback : str = None):
    """
    Extracts the title from a DOCX file.

    This function attempts the following in order:
    1. Finds the first paragraph styled as a heading or containing bold text.
    2. Uses a fallback if provided.
    3. Falls back to the filename as the title.

    Parameters:
        docx_path (str): Path to the DOCX file.
        fallback (str, optional): Fallback title if no heading or bold text is found.

    Returns:
        str: The extracted title.
    """
    doc = Document(docx_path)
    
    # Try to find a bold or large font text as the title
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") or any(run.bold for run in para.runs):
            return para.text.strip()
    if fallback:
        return fallback
    # If no heading found, assume first paragraph is the title
    return os.path.basename(docx_path)


def extract_text_from_pdf(pdf_path):
    """
    Extracts all text from a PDF file.

    Parameters:
        pdf_path (str): Path to the PDF file.

    Returns:
        str: Combined text content from all pages.
    """
    text = ""
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    """
    Extracts all text from a DOCX file.

    Parameters:
        docx_path (str): Path to the DOCX file.

    Returns:
        str: Combined text content from all paragraphs.
    """
    doc = Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs])